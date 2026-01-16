#!/usr/bin/env python3
"""Convert markdown files to Word with proper formatting."""

import re
import subprocess
import os
import sys
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import qn

def set_table_borders(table):
    """Add black borders to a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(r'<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')

    tblBorders = parse_xml(
        r'<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        r'<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        r'<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        r'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        r'<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        r'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        r'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        r'</w:tblBorders>'
    )

    tblPr.append(tblBorders)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)

def set_table_full_width(table):
    """Set table width to 100% of page width."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(r'<w:tblPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
        tbl.insert(0, tblPr)

    # Remove existing width setting if any
    for child in list(tblPr):
        if child.tag.endswith('}tblW'):
            tblPr.remove(child)

    # Set width to 100% (5000 = 100% in twentieths of a percent)
    tblW = parse_xml(r'<w:tblW xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:type="pct" w:w="5000"/>')
    tblPr.insert(0, tblW)

def escape_xml(text):
    """Escape special XML characters."""
    text = text.replace('&', '&amp;')
    text = text.replace('<', '&lt;')
    text = text.replace('>', '&gt;')
    text = text.replace('"', '&quot;')
    text = text.replace("'", '&apos;')
    return text

def convert_markers_to_linebreaks_xml(cell):
    """Convert ⏎ markers to actual line breaks in a cell, preserving hyperlinks."""
    full_text = cell.text

    if '⏎' not in full_text:
        return False

    tc = cell._tc
    text_elements = tc.findall('.//' + qn('w:t'))

    for t_elem in text_elements:
        if t_elem.text and '⏎' in t_elem.text:
            parts = t_elem.text.split('⏎')
            if len(parts) > 1:
                t_elem.text = parts[0].strip()
                insert_after = t_elem
                for part in parts[1:]:
                    part = part.strip()
                    br = parse_xml(r'<w:br xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
                    insert_after.addnext(br)
                    insert_after = br
                    if part:
                        escaped_part = escape_xml(part)
                        new_t = parse_xml(f'<w:t xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xml:space="preserve">{escaped_part}</w:t>')
                        insert_after.addnext(new_t)
                        insert_after = new_t
    return True

def preprocess_markdown(content):
    """Preprocess markdown content for Word conversion."""
    # Remove the div wrapper for black-border-table
    content = re.sub(r'<div class="black-border-table" markdown="1">\s*\n?', '', content)
    content = re.sub(r'\n?</div>', '', content)

    # Remove Jekyll front matter
    content = re.sub(r'^---\n.*?\n---\n', '', content, flags=re.DOTALL)

    # Remove navigation links at top
    content = re.sub(r'^\[.*?\]\(index\).*?\n+', '', content)
    content = re.sub(r'^\*\*\[Download Word Document\].*?\*\*\n+', '', content)

    # Convert relative links to absolute GitHub Pages URLs
    base_url = "https://sunjinpak.github.io/MGMT6100/"

    def convert_link(match):
        text = match.group(1)
        url = match.group(2)
        if url.startswith('http') or url.startswith('#') or url.startswith('mailto:'):
            return match.group(0)
        if url.endswith('.docx'):
            return match.group(0)
        return f'[{text}]({base_url}{url})'

    content = re.sub(r'\[([^\]]+)\]\(([^)]+)\)', convert_link, content)

    # Replace <br> with marker
    content = content.replace('<br>', '⏎')

    # Ensure bullet lists have a blank line before them
    content = re.sub(r'([^\n])\n(- )', r'\1\n\n\2', content)

    # Ensure tables have a blank line before them (heading followed by table)
    content = re.sub(r'(##[^\n]*)\n(\|)', r'\1\n\n\2', content)

    # Remove dashes in table cells
    lines = content.split('\n')
    processed_lines = []
    for line in lines:
        if line.startswith('|') and '-----' in line and not re.match(r'^\|[-:\s|]+\|$', line):
            line = re.sub(r'-{5,}', '', line)
        processed_lines.append(line)
    content = '\n'.join(processed_lines)

    return content

def add_bookmark_hyperlink(paragraph, text, bookmark_name):
    """Add a hyperlink to a bookmark within the document."""
    # Create hyperlink element with anchor (internal bookmark reference)
    hyperlink = parse_xml(
        f'<w:hyperlink xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:anchor="{bookmark_name}">'
        f'<w:r><w:rPr><w:color w:val="0066CC"/><w:u w:val="single"/></w:rPr>'
        f'<w:t>{escape_xml(text)}</w:t></w:r></w:hyperlink>'
    )
    paragraph._p.append(hyperlink)

def add_toc_links(doc):
    """Add internal bookmark links to Table of Contents."""
    # TOC mapping: display text -> bookmark name
    toc_items = [
        ("Course Description", "course-description"),
        ("BSBA Program Learning Objectives", "bsba-program-learning-objectives"),
        ("Course Learning Objectives", "course-learning-objectives"),
        ("Required Text and Materials", "required-text-and-materials"),
        ("Instructional Strategies & Classroom Policies", "instructional-strategies-classroom-policies"),
        ("Course Assessment", "course-assessment"),
        ("Course Schedule", "course-schedule"),
        ("CSUB Student Chapter – SHRM", "csub-student-chapter-shrm"),
        ("University Policies", "university-policies"),
    ]

    # Find TOC section and replace text with hyperlinks
    in_toc = False
    toc_count = 0
    for para in doc.paragraphs:
        if "Table of Contents" in para.text:
            in_toc = True
            continue

        if in_toc and toc_count < len(toc_items):
            # Check if this paragraph matches a TOC item
            para_text = para.text.strip()
            for display_text, bookmark in toc_items:
                if display_text == para_text or para_text.endswith(display_text):
                    # Clear paragraph XML and rebuild with hyperlink
                    p = para._p
                    # Remove all runs
                    for child in list(p):
                        if child.tag.endswith('}r') or child.tag.endswith('}hyperlink'):
                            p.remove(child)
                    # Add hyperlink
                    add_bookmark_hyperlink(para, display_text, bookmark)
                    toc_count += 1
                    break

            # Stop if we've processed all TOC items or hit next section
            if para.style and para.style.name.startswith('Heading') and "Table of Contents" not in para.text:
                break

def postprocess_word(doc, filename=''):
    """Post-process Word document."""
    # Add TOC internal links for Syllabus
    if 'syllabus' in filename.lower():
        add_toc_links(doc)

    # Add borders and set full width for all tables
    for table in doc.tables:
        set_table_borders(table)
        set_table_full_width(table)

    # Process tables for line break markers
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                convert_markers_to_linebreaks_xml(cell)

    # Set column widths for specific tables (Syllabus)
    if 'syllabus' in filename.lower():
        for table in doc.tables:
            first_cell = table.rows[0].cells[0].text.strip().lower()

            # Weekly Schedule table (2 columns: Class, Contents)
            if len(table.columns) == 2 and 'class' in first_cell:
                for row in table.rows:
                    row.cells[0].width = Cm(3)
                    row.cells[1].width = Cm(14)

            # Grading Criteria table (3 columns: Assessment, Points, %)
            if len(table.columns) == 3 and 'assessment' in first_cell:
                for row in table.rows:
                    row.cells[0].width = Cm(9)
                    row.cells[1].width = Cm(5)
                    row.cells[2].width = Cm(2)

            # Course Structure table (3 columns: Stage, Module, Week)
            if len(table.columns) == 3 and 'stage' in first_cell:
                for row in table.rows:
                    row.cells[0].width = Cm(4)
                    row.cells[1].width = Cm(9)
                    row.cells[2].width = Cm(2)

            # Grading Scale table (2 columns: Grade, Range)
            if len(table.columns) == 2 and 'grade' in first_cell:
                for row in table.rows:
                    row.cells[0].width = Cm(3)
                    row.cells[1].width = Cm(5)

            # Important Dates table (2 columns: Date, Event)
            if len(table.columns) == 2 and 'date' in first_cell:
                for row in table.rows:
                    row.cells[0].width = Cm(4)
                    row.cells[1].width = Cm(12)

    # Make all hyperlinks blue and underlined
    body = doc._body._body
    hyperlinks = body.findall('.//' + qn('w:hyperlink'))

    for hyperlink in hyperlinks:
        runs = hyperlink.findall('.//' + qn('w:r'))
        for run_elem in runs:
            rPr = run_elem.find(qn('w:rPr'))
            if rPr is None:
                rPr = parse_xml(r'<w:rPr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
                run_elem.insert(0, rPr)

            color = rPr.find(qn('w:color'))
            if color is None:
                color = parse_xml(r'<w:color xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="0066CC"/>')
                rPr.append(color)
            else:
                color.set(qn('w:val'), '0066CC')

            u = rPr.find(qn('w:u'))
            if u is None:
                u = parse_xml(r'<w:u xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="single"/>')
                rPr.append(u)

def convert_md_to_word(md_file, docx_file):
    """Convert a markdown file to Word."""
    print(f"Converting {md_file} -> {docx_file}")

    # Read and preprocess markdown
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()

    content = preprocess_markdown(content)

    # Write temp file
    temp_file = 'temp_convert.md'
    with open(temp_file, 'w', encoding='utf-8') as f:
        f.write(content)

    # Convert with pandoc
    subprocess.run([
        'pandoc', temp_file,
        '-o', docx_file,
        '--from=markdown+raw_html'
    ], check=True)

    # Post-process
    doc = Document(docx_file)
    postprocess_word(doc, docx_file)
    doc.save(docx_file)

    # Cleanup
    os.remove(temp_file)
    print(f"  Created: {docx_file}")

def fix_markdown_tables(md_file):
    """Ensure tables in markdown have blank lines before them."""
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()

    original = content

    # Ensure tables have a blank line before them (after headings)
    content = re.sub(r'(##[^\n]*)\n(\|)', r'\1\n\n\2', content)

    # Ensure tables have a blank line before them (after text ending with :)
    content = re.sub(r'([^\n|])\n(\|[^\-])', r'\1\n\n\2', content)

    if content != original:
        with open(md_file, 'w', encoding='utf-8') as f:
            f.write(content)
        print(f"  Fixed table formatting in {md_file}")
        return True
    return False

def main():
    # Get the directory where this script is located
    script_dir = os.path.dirname(os.path.abspath(__file__))
    os.chdir(script_dir)

    # Files to process (excluding index.md and README.md)
    md_files = [
        ('syllabus.md', 'Syllabus.docx'),
        ('roster.md', 'Roster.docx'),
        ('team-meeting-rubric.md', 'Team-Meeting-Rubric.docx'),
        ('team-presentation-rubric.md', 'Team-Presentation-Rubric.docx'),
        ('team-meeting-1.md', 'Team-Meeting-1.docx'),
        ('team-meeting-2.md', 'Team-Meeting-2.docx'),
        ('team-meeting-3.md', 'Team-Meeting-3.docx'),
        ('team-meeting-4.md', 'Team-Meeting-4.docx'),
        ('team-meeting-5.md', 'Team-Meeting-5.docx'),
        ('team-meeting-6.md', 'Team-Meeting-6.docx'),
        ('team-meeting-7.md', 'Team-Meeting-7.docx'),
        ('team-meeting-8.md', 'Team-Meeting-8.docx'),
        ('team-meeting-9.md', 'Team-Meeting-9.docx'),
        ('team-meeting-10.md', 'Team-Meeting-10.docx'),
        ('team-meeting-11.md', 'Team-Meeting-11.docx'),
    ]

    print("Step 1: Fixing markdown table formatting...")
    for md_file, _ in md_files:
        if os.path.exists(md_file):
            fix_markdown_tables(md_file)

    print("\nStep 2: Converting to Word...")
    for md_file, docx_file in md_files:
        if os.path.exists(md_file):
            convert_md_to_word(md_file, docx_file)

    print("\nDone! All files converted.")

if __name__ == "__main__":
    main()
